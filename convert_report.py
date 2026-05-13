"""Convert project progress report from Markdown to Word .docx"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# ---- Style setup ----
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)

# Heading styles
for level, (size, bold) in enumerate([(22, True), (16, True), (13, True), (12, True)], 1):
    h_style = doc.styles[f'Heading {level}']
    h_font = h_style.font
    h_font.name = '黑体'
    h_font.size = Pt(size)
    h_font.bold = bold
    h_font.color.rgb = RGBColor(0, 0, 0)
    h_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def add_paragraph(text, bold=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    return p

def add_code_block(code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = '宋体'
        run.font.size = Pt(10)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = '宋体'
            run.font.size = Pt(10)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # spacing after table
    return table

def parse_and_build(md_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i].rstrip()

        # Code block toggle
        if line.startswith('```'):
            if in_code_block:
                add_code_block('\n'.join(code_lines))
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Separator
        if line.startswith('---'):
            doc.add_paragraph()
            i += 1
            continue

        # Headings
        if line.startswith('# ') and not line.startswith('## '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)

        # Table detection: lines starting with |
        elif line.startswith('|') and line.endswith('|'):
            table_lines.append(line)
            # Check if next line is separator or next row
            if i + 1 < len(lines) and lines[i + 1].strip().startswith('|'):
                i += 1
                continue
            elif table_lines:
                # Process collected table
                rows_data = []
                headers_data = []
                for ti, tl in enumerate(table_lines):
                    cells = [c.strip() for c in tl.split('|')[1:-1]]
                    if ti == 0:
                        headers_data = cells
                    elif not all(re.match(r'^[-: ]+$', c) for c in cells):
                        rows_data.append(cells)
                if headers_data:
                    add_table(headers_data, rows_data)
                table_lines = []
            i += 1
            continue

        # Unordered list
        elif line.startswith('- ') or line.startswith('  - '):
            text = re.sub(r'^\s*-\s+', '', line)
            # Bold prefix up to first colon
            p = doc.add_paragraph(style='List Bullet')
            text = re.sub(r'\*\*(.+?)\*\*', lambda m: m.group(1), text)
            text = re.sub(r'`(.+?)`', lambda m: m.group(1), text)
            run = p.add_run(text)
            run.font.name = '宋体'
            run.font.size = Pt(11)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # Ordered list
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            text = re.sub(r'\*\*(.+?)\*\*', lambda m: m.group(1), text)
            text = re.sub(r'`(.+?)`', lambda m: m.group(1), text)
            run = p.add_run(text)
            run.font.name = '宋体'
            run.font.size = Pt(11)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # Bold-only line like **xxx**
        elif line.startswith('**') and line.endswith('**'):
            add_paragraph(line.strip('*'), bold=True)

        # Empty line
        elif not line.strip():
            pass

        # Regular paragraph
        else:
            # Strip markdown formatting
            text = line
            text = re.sub(r'\*\*(.+?)\*\*', lambda m: m.group(1), text)
            text = re.sub(r'`(.+?)`', lambda m: m.group(1), text)
            text = re.sub(r'\[(.+?)\]\(.+?\)', lambda m: m.group(1), text)
            if text.strip():
                add_paragraph(text, indent=True)

        i += 1

# Build document
md_file = r'e:\智慧水利应用\智慧水利应用-功能开发与项目进度报告.md'
parse_and_build(md_file)

# Set landscape for wide tables
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

output_path = r'e:\智慧水利应用\智慧水利应用-功能开发与项目进度报告.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
